import os
import shutil
import json
from docx import Document
from PyPDF2 import PdfReader
from transformers import pipeline
import pytesseract
from pdf2image import convert_from_path
from concurrent.futures import ThreadPoolExecutor
import argparse

# Load the summarization model (BART)
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# Function to copy files and validate integrity
def copy_files_to_directory(source_dir, data_dir, file_names):
    os.makedirs(data_dir, exist_ok=True)
    for file_name in file_names:
        source_path = os.path.join(source_dir, file_name)
        destination_path = os.path.join(data_dir, file_name)
        
        if not os.path.exists(source_path):
            print(f"⚠️ File not found: {source_path}")
            continue
        
        try:
            shutil.copy(source_path, destination_path)
            print(f"✅ Copied {source_path} to {destination_path}")
            
            # Validate file integrity after copying
            if file_name.endswith(".pdf"):
                with open(destination_path, "rb") as f:
                    PdfReader(f)
            elif file_name.endswith(".docx"):
                Document(destination_path) 
                
        except Exception as e:
            print(f"❌ Corrupted file detected: {file_name} (Error: {e})")
            os.remove(destination_path) 

# Function to read file content
def read_file(filepath):
    try:
        if filepath.endswith(".txt"):
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()
        
        elif filepath.endswith(".docx"):
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
            if not text.strip():
                print(f"⚠️ DOCX file has no readable text: {filepath}")
            return text
        
        elif filepath.endswith(".pdf"):
            with open(filepath, "rb") as f:
                pdf_reader = PdfReader(f)
                text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
                if not text.strip():
                    print(f"⚠️ PDF has no extractable text (may be scanned/image-based): {filepath}")
                    # Use OCR as a fallback
                    images = convert_from_path(filepath)
                    text = "\n".join([pytesseract.image_to_string(image) for image in images])
                return text
        
        elif filepath.endswith(".json"):
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("content", json.dumps(data, indent=2))
    
    except Exception as e:
        print(f"❌ Failed to read {filepath}: {e}")
    return ""

# Function to generate a summary
def generate_summary(text):
    if not text.strip():
        return "No text to summarize."
    
    # Split text into chunks of 1024 tokens
    max_chunk_size = 1024
    text_chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    
    summaries = []
    for chunk in text_chunks:
        try:
            summary = summarizer(chunk, max_length=100, min_length=30, do_sample=False)
            summaries.append(summary[0]['summary_text'])
        except Exception as e:
            return f"Summarization failed: {str(e)}"
    
    return " ".join(summaries)

# Function to process a single file
def process_file(filepath, data_dir):
    content = read_file(filepath)
    
    if not content:
        print(f"⚠️ Skipping corrupted/unreadable file: {filepath}")
        return
    
    print(f"\nProcessing: {filepath}")
    print(f"Content length: {len(content.split())} words")
    
    summary = generate_summary(content)
    summary_path = os.path.join(data_dir, os.path.basename(filepath) + "_summary.txt")
    
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(summary)
    print(f"✅ Summary saved: {summary_path}")

# Main workflow
def main(file_names):
   
    source_dir = os.getcwd()
    data_dir = os.path.join(source_dir, "processed_files")  

    
    os.makedirs(data_dir, exist_ok=True)

    print("\n--- Copying Files to Processed Directory ---")
    copy_files_to_directory(source_dir, data_dir, file_names)

    print("\n--- Summarization and Saving Summaries ---")
    with ThreadPoolExecutor() as executor:
        for root, _, files in os.walk(data_dir):
            filepaths = [os.path.join(root, file) for file in files]
            executor.map(lambda filepath: process_file(filepath, data_dir), filepaths)

    print("\n--- Process completed ---")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Summarize text from various file formats.")
    parser.add_argument("file_names", nargs='+', help="List of file names to process.")
    
    args = parser.parse_args()
    main(args.file_names)
